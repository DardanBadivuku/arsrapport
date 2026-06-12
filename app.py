from flask import Flask, request, send_file, jsonify
import os, io, zipfile, tempfile, subprocess, shutil
import pandas as pd
from docx import Document

app = Flask(__name__, static_folder='.', template_folder='.')
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024

BASE = os.path.dirname(os.path.abspath(__file__))

@app.route('/')
def index():
    return open(os.path.join(BASE, 'index.html'), encoding='utf-8').read()

def find_col(df, candidates):
    for c in candidates:
        for col in df.columns:
            if c.lower() in col.lower():
                return col
    return None

def replace_para(para, old, new):
    if old not in para.text: return
    new_text = para.text.replace(old, new)
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]: r.text = ''

def update_doc(doc, navn, adresse, postnr, poststed, kontakt, yf, yt):
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt == 'Kunde':
            if para.runs: para.runs[0].text = navn; [setattr(r,'text','') for r in para.runs[1:]]
        elif txt == 'Adresse….':
            if para.runs: para.runs[0].text = adresse; [setattr(r,'text','') for r in para.runs[1:]]
        elif txt == 'Post/sted':
            ps = f'{postnr} {poststed}'.strip()
            if para.runs: para.runs[0].text = ps; [setattr(r,'text','') for r in para.runs[1:]]
        elif txt.startswith('V/') and len(txt) <= 4:
            new_v = f'V/ {kontakt}' if kontakt else 'V/'
            if para.runs: para.runs[0].text = new_v; [setattr(r,'text','') for r in para.runs[1:]]
        if yf in para.text:
            replace_para(para, yf, yt)
    for section in doc.sections:
        for hf in [section.header, section.footer]:
            if hf:
                for p in hf.paragraphs:
                    if yf in p.text: replace_para(p, yf, yt)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if yf in p.text: replace_para(p, yf, yt)

@app.route('/preview', methods=['POST'])
def preview():
    try:
        excel_file = request.files.get('excel')
        if not excel_file:
            return jsonify({'error': 'Ingen fil'}), 400
        df = pd.read_excel(excel_file, sheet_name=0)
        df.columns = [str(c).strip() for c in df.columns]
        navn_col    = find_col(df, ['Navn','Name'])
        kontakt_col = find_col(df, ['Kontaktperson','Kontakt','Contact'])
        if not navn_col:
            return jsonify({'error': 'Finner ikke kolonnen "Navn" i Excel-filen'}), 400
        df = df.dropna(subset=[navn_col])
        kunder = []
        for _, row in df.head(10).iterrows():
            navn = str(row.get(navn_col,'')).strip()
            k    = str(row.get(kontakt_col,'')).strip() if kontakt_col else ''
            if navn and navn != 'nan':
                kunder.append({'navn': navn, 'kontakt': '' if k == 'nan' else k})
        return jsonify({'kunder': kunder, 'totalt': len(df)})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/generate', methods=['POST'])
def generate():
    try:
        word_file  = request.files.get('word')
        excel_file = request.files.get('excel')
        yr_from    = request.form.get('year_from', '2024')
        yr_to      = request.form.get('year_to',   '2026')

        if not word_file or not excel_file:
            return jsonify({'error': 'Mangler filer'}), 400

        df = pd.read_excel(excel_file, sheet_name=0)
        df.columns = [str(c).strip() for c in df.columns]

        navn_col    = find_col(df, ['Navn','Name'])
        adr_col     = find_col(df, ['Adresse','Address'])
        postnr_col  = find_col(df, ['Postnr','Postkode','Zip'])
        poststed_col= find_col(df, ['Poststed','Sted','City'])
        kontakt_col = find_col(df, ['Kontaktperson','Kontakt','Contact'])

        if not navn_col:
            return jsonify({'error': 'Finner ikke kolonnen "Navn" i Excel-filen'}), 400

        df = df.dropna(subset=[navn_col])
        word_bytes = word_file.read()

        def clean(val):
            s = str(val).strip()
            return '' if s in ('nan','None','<NA>') else s

        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
            for _, row in df.iterrows():
                navn = clean(row.get(navn_col, ''))
                if not navn: continue

                adresse  = clean(row.get(adr_col,''))     if adr_col else ''
                raw_pnr  = clean(row.get(postnr_col,''))  if postnr_col else ''
                postnr   = raw_pnr.split('.')[0] if raw_pnr else ''
                poststed = clean(row.get(poststed_col,'')) if poststed_col else ''
                kontakt  = clean(row.get(kontakt_col,''))  if kontakt_col else ''

                try:
                    doc = Document(io.BytesIO(word_bytes))
                    update_doc(doc, navn, adresse, postnr, poststed, kontakt, yr_from, yr_to)

                    tmp = tempfile.mkdtemp()
                    tmp_docx = os.path.join(tmp, 'rapport.docx')
                    doc.save(tmp_docx)

                    subprocess.run(
                        ['libreoffice','--headless','--convert-to','pdf','--outdir',tmp,tmp_docx],
                        capture_output=True, timeout=90
                    )
                    tmp_pdf = os.path.join(tmp, 'rapport.pdf')
                    safe = ''.join(c for c in navn if c not in r'\/:*?"<>|').strip()
                    pdf_name = f'Årsrapport {yr_to} - {safe}.pdf'

                    if os.path.exists(tmp_pdf):
                        zf.write(tmp_pdf, pdf_name)

                    shutil.rmtree(tmp, ignore_errors=True)
                except Exception:
                    shutil.rmtree(tmp, ignore_errors=True) if 'tmp' in dir() else None

        zip_buf.seek(0)
        return send_file(
            zip_buf,
            mimetype='application/zip',
            as_attachment=True,
            download_name=f'Årsrapporter_{yr_to}.zip'
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8080))
    app.run(host='0.0.0.0', port=port)
